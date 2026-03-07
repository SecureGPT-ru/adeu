"""
Тесты для EditResult — детальной обратной связи о применении правок.

Проверяет:
- get_edit_results() возвращает список EditResult
- EditResult содержит правильный status (APPLIED, SKIPPED_NOT_FOUND, SKIPPED_OVERLAP)
- matched_text содержит фактически найденный текст (fuzzy matching)
"""

import io

from docx import Document

from adeu.models import DocumentEdit, EditResult, EditStatus
from adeu.redline.engine import RedlineEngine


def test_edit_result_applied_basic():
    """Базовый тест: успешная правка возвращает status=APPLIED и matched_text."""
    doc = Document()
    doc.add_paragraph("Hello World")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    edit = DocumentEdit(target_text="Hello", new_text="Hi")

    engine = RedlineEngine(stream)
    applied, skipped = engine.apply_edits([edit])

    assert applied == 1
    assert skipped == 0

    results = engine.get_edit_results()
    assert len(results) == 1

    result = results[0]
    assert result.status == EditStatus.APPLIED
    assert result.target_text == "Hello"
    assert result.new_text == "Hi"
    assert result.matched_text == "Hello"


def test_edit_result_not_found():
    """Правка с несуществующим target_text возвращает status=SKIPPED_NOT_FOUND."""
    doc = Document()
    doc.add_paragraph("Hello World")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    edit = DocumentEdit(target_text="Goodbye", new_text="Hi")

    engine = RedlineEngine(stream)
    applied, skipped = engine.apply_edits([edit])

    assert applied == 0
    assert skipped == 1

    results = engine.get_edit_results()
    assert len(results) == 1

    result = results[0]
    assert result.status == EditStatus.SKIPPED_NOT_FOUND
    assert result.target_text == "Goodbye"
    assert result.matched_text is None


def test_edit_result_overlap():
    """Перекрывающиеся правки: вторая получает status=SKIPPED_OVERLAP."""
    doc = Document()
    doc.add_paragraph("Hello World")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Обе правки пытаются изменить "Hello"
    edit1 = DocumentEdit(target_text="Hello", new_text="Hi")
    edit2 = DocumentEdit(target_text="Hello World", new_text="Greetings")

    engine = RedlineEngine(stream)
    applied, skipped = engine.apply_edits([edit1, edit2])

    # edit1 применится, edit2 будет пропущен из-за перекрытия
    assert applied == 1
    assert skipped == 1

    results = engine.get_edit_results()
    assert len(results) == 2

    # Первая правка применена
    assert results[0].status == EditStatus.APPLIED
    # Вторая пропущена из-за перекрытия
    assert results[1].status == EditStatus.SKIPPED_OVERLAP


def test_edit_result_fuzzy_matching():
    """Fuzzy matching: matched_text отличается от target_text."""
    doc = Document()
    # Документ содержит 10 подчёркиваний
    doc.add_paragraph("Sign here: [__________]")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Ищем только 3 подчёркивания — fuzzy matching должен найти 10
    edit = DocumentEdit(target_text="Sign here: [___]", new_text="Sign here: John Doe")

    engine = RedlineEngine(stream)
    applied, skipped = engine.apply_edits([edit])

    assert applied == 1

    results = engine.get_edit_results()
    result = results[0]

    assert result.status == EditStatus.APPLIED
    # matched_text содержит фактически найденный текст (10 подчёркиваний)
    assert "[__________]" in result.matched_text


def test_edit_result_multiple_edits():
    """Несколько правок: каждая получает свой EditResult."""
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    doc.add_paragraph("Third paragraph.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    edits = [
        DocumentEdit(target_text="First", new_text="1st"),
        DocumentEdit(target_text="Second", new_text="2nd"),
        DocumentEdit(target_text="Fourth", new_text="4th"),  # Не существует
    ]

    engine = RedlineEngine(stream)
    applied, skipped = engine.apply_edits(edits)

    assert applied == 2
    assert skipped == 1

    results = engine.get_edit_results()
    assert len(results) == 3

    # Проверяем результаты по содержимому, т.к. порядок может отличаться из-за сортировки
    results_by_target = {r.target_text: r for r in results}

    assert results_by_target["First"].status == EditStatus.APPLIED
    assert results_by_target["Second"].status == EditStatus.APPLIED
    assert results_by_target["Fourth"].status == EditStatus.SKIPPED_NOT_FOUND


def test_get_edit_results_empty_before_apply():
    """get_edit_results() возвращает пустой список до вызова apply_edits()."""
    doc = Document()
    doc.add_paragraph("Hello")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream)
    results = engine.get_edit_results()

    assert results == []


def test_edit_result_with_comment():
    """EditResult сохраняет комментарий из исходного DocumentEdit."""
    doc = Document()
    doc.add_paragraph("Contract text here.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    edit = DocumentEdit(
        target_text="Contract",
        new_text="Agreement",
        comment="Заменяем 'Contract' на 'Agreement' для единообразия"
    )

    engine = RedlineEngine(stream)
    engine.apply_edits([edit])

    results = engine.get_edit_results()
    assert len(results) == 1

    result = results[0]
    assert result.status == EditStatus.APPLIED
    assert result.comment == "Заменяем 'Contract' на 'Agreement' для единообразия"
